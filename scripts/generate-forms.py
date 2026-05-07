#!/usr/bin/env python3
"""Generate Audrey Baliao's intake form and service contract as fillable
Word documents. Output goes into /public so the live site can serve them
directly via /<filename>.docx.

Run:    python3 scripts/generate-forms.py
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    print("Installing python-docx (one-time)...", flush=True)
    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "--quiet", "python-docx"]
    )
    from docx import Document  # noqa: E402
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
    from docx.oxml import OxmlElement  # noqa: E402
    from docx.oxml.ns import qn  # noqa: E402
    from docx.shared import Cm, Pt, RGBColor  # noqa: E402


ROOT = Path(__file__).resolve().parent.parent
PUBLIC = ROOT / "public"
PUBLIC.mkdir(exist_ok=True)

# Brand palette: matches the website's Tailwind emerald/gold tokens.
EMERALD_950 = RGBColor(0x0A, 0x24, 0x18)
EMERALD_700 = RGBColor(0x1A, 0x61, 0x3E)
EMERALD_500 = RGBColor(0x2C, 0x94, 0x5F)
GOLD_500 = RGBColor(0xD8, 0xA8, 0x3A)
GOLD_600 = RGBColor(0xB8, 0x8A, 0x26)
INK = RGBColor(0x14, 0x3F, 0x2B)
MUTED = RGBColor(0x60, 0x6B, 0x60)

BLANK_LONG = "_" * 70
BLANK_MED = "_" * 40
BLANK_SHORT = "_" * 22


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def style_run(run, *, size=11, bold=False, italic=False, color=INK):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if level == 1:
        style_run(r, size=20, bold=True, color=EMERALD_950)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
    elif level == 2:
        style_run(r, size=13, bold=True, color=EMERALD_700)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
    else:
        style_run(r, size=11, bold=True, color=EMERALD_700)


def body(doc, text, *, italic=False, color=None, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size=size, italic=italic, color=color or INK)
    return p


def field(doc, label, blank=BLANK_LONG, hint=""):
    """A 'fill-in' line: bold label, then an underlined blank, optional hint."""
    p = doc.add_paragraph()
    r1 = p.add_run(label + ":  ")
    style_run(r1, size=10.5, bold=True, color=EMERALD_950)
    r2 = p.add_run(blank)
    style_run(r2, size=10.5, color=MUTED)
    if hint:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        r3 = p2.add_run(hint)
        style_run(r3, size=9, italic=True, color=MUTED)


def add_accent_rule(doc, color="d8a83a", size_pt=8):
    """Thin gold horizontal rule below the letterhead."""
    accent = doc.add_paragraph()
    pPr = accent._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    accent.paragraph_format.space_after = Pt(8)


def add_letterhead(doc):
    """2-column header: name on the left, contact details right-aligned."""
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(10)
    t.columns[1].width = Cm(7)

    left = t.rows[0].cells[0]
    p1 = left.paragraphs[0]
    r1 = p1.add_run("Dhey Creates")
    style_run(r1, size=24, bold=True, color=EMERALD_950)
    p2 = left.add_paragraph()
    r2 = p2.add_run("Audrey Baliao · Video Editor")
    style_run(r2, size=10.5, color=EMERALD_700)

    right = t.rows[0].cells[1]
    contact_lines = [
        "audreybaliao022@gmail.com",
        "Instagram · @ur.dhey",
        "TikTok · @.dtb8 / @adryrzbl",
        "Facebook · audrey.baliao.2024",
        "Philippines",
    ]
    for i, line in enumerate(contact_lines):
        p = right.paragraphs[0] if i == 0 else right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(line)
        style_run(r, size=9.5, color=MUTED)

    add_accent_rule(doc)


def page_setup(doc):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)


def fill_lines(doc, n=3):
    """N blank underlined paragraphs, for free-form text fields."""
    for _ in range(n):
        p = doc.add_paragraph()
        r = p.add_run(BLANK_LONG + BLANK_LONG[:30])
        style_run(r, size=10.5, color=MUTED)


# ---------------------------------------------------------------------------
# Intake form
# ---------------------------------------------------------------------------
def build_intake_form() -> Path:
    doc = Document()
    page_setup(doc)
    add_letterhead(doc)

    heading(doc, "Project intake form")
    body(
        doc,
        "Use this form before I quote. Fill in what you know; anything you "
        "are unsure of, leave blank and we will cover it on the discovery "
        "call.",
        italic=True,
        color=MUTED,
    )
    doc.add_paragraph()

    heading(doc, "Your details", level=2)
    field(doc, "Name", BLANK_LONG)
    field(
        doc,
        "Company / brand",
        BLANK_LONG,
        hint="Optional. Only if you are booking on behalf of a brand.",
    )
    field(doc, "Email", BLANK_LONG)
    field(doc, "Phone or messaging app", BLANK_MED)
    field(
        doc,
        "Best way to reach you",
        BLANK_MED,
        hint="Email, WhatsApp, Instagram DM, Messenger, etc.",
    )

    heading(doc, "Project basics", level=2)
    field(doc, "Project title", BLANK_LONG)
    body(
        doc,
        "Project type (tick one):",
    )
    body(
        doc,
        "  ☐  Vlog edit          ☐  Music video edit          "
        "☐  News / project edit",
    )
    body(
        doc,
        "  ☐  Travel / wishlist edit          ☐  Photo-dump carousel          "
        "☐  Other",
    )
    field(
        doc,
        "Platform",
        BLANK_MED,
        hint="YouTube / Facebook / Instagram / TikTok / other.",
    )
    field(doc, "Target post date", BLANK_MED)

    heading(doc, "Footage", level=2)
    field(doc, "Raw footage length (minutes)", BLANK_SHORT)
    field(
        doc,
        "Google Drive link to footage",
        BLANK_LONG,
        hint=(
            'Set sharing to "Anyone with the link can view" so I can access '
            "it without delays."
        ),
    )
    body(doc, "What is the video about? (2 to 3 sentences)")
    fill_lines(doc, 3)

    heading(doc, "Creative direction", level=2)
    field(
        doc,
        "Vibe / tone",
        BLANK_LONG,
        hint="Cinematic / energetic / chill / professional / etc.",
    )
    field(
        doc,
        "Reference videos or styles",
        BLANK_LONG,
        hint="Optional. Paste 1 to 3 links.",
    )

    heading(doc, "Edit choices", level=2)
    body(doc, "Tick the level you want for each. Pricing scales with complexity.")
    body(doc, "  Music & sound design:    ☐ none   ☐ Light   ☐ Standard   ☐ Detailed   ☐ Heavy")
    body(doc, "  Video effects:                  ☐ none   ☐ Light   ☐ Standard   ☐ Detailed   ☐ Heavy")
    body(doc, "  Subtitles / captions:        ☐ none   ☐ Light   ☐ Standard   ☐ Detailed   ☐ Heavy")
    doc.add_paragraph()
    body(doc, "Hooks & branding (tick what you want):")
    body(doc, "  ☐ Cliffhanger opener   ☐ Custom intro   ☐ Custom outro")
    body(doc, "  ☐ Thumbnail   ☐ Cover frame")
    body(doc, "  ☐ YouTube kit (timestamps + optimized title + description, ₱100 all-in)")

    heading(doc, "Delivery", level=2)
    field(
        doc,
        "Extra export versions",
        BLANK_SHORT,
        hint="Vertical 9:16, square 1:1, cutdowns. Count beyond the main deliverable.",
    )
    body(doc, "Delivery speed:   ☐ Standard (7-day)   ☐ Rush (5-day, +25%)")

    heading(doc, "Discount request (optional)", level=2)
    body(doc, "  ☐ 5%   paying the full month upfront instead of 50/50")
    body(doc, "  ☐ 10%  3-month commitment")
    body(doc, "  ☐ 15%  6-month commitment, or referring another paying creator")
    body(doc, "  ☐ 20%  12-month commitment, or multi-creator package (cap)")
    field(
        doc,
        "Or write your own reasoning",
        BLANK_LONG,
        hint="Reasonable counter-offers welcome.",
    )

    heading(doc, "Budget & notes", level=2)
    field(
        doc,
        "Budget range (₱)",
        BLANK_MED,
        hint="Optional. Helps me match scope to your budget.",
    )
    body(doc, "Anything else you would like me to know:")
    fill_lines(doc, 3)

    out = PUBLIC / "dhey-creates-intake-form.docx"
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# Service contract
# ---------------------------------------------------------------------------
def build_service_contract() -> Path:
    doc = Document()
    page_setup(doc)
    add_letterhead(doc)

    heading(doc, "Service contract for video editing")
    body(
        doc,
        "Signed once a project is booked. Fill in the blanks, sign at the "
        "bottom, and return a scanned or photographed copy.",
        italic=True,
        color=MUTED,
    )
    doc.add_paragraph()

    heading(doc, "1. Parties", level=2)
    body(doc, "This agreement is made between:")
    field(
        doc,
        "  Editor",
        "Audrey Baliao, Philippines",
    )
    field(doc, "  Client", BLANK_LONG)
    field(doc, "  Client address / location", BLANK_LONG)
    field(doc, "  Date of agreement", BLANK_MED)

    heading(doc, "2. Project scope", level=2)
    field(doc, "Project title", BLANK_LONG)
    body(doc, "Scope summary (what I will edit and deliver):")
    fill_lines(doc, 3)
    body(
        doc,
        "The full creative brief is captured in the project intake form, "
        "which forms part of this agreement by reference.",
        italic=True,
        color=MUTED,
    )

    heading(doc, "3. Deliverables", level=2)
    body(doc, "I will deliver:")
    body(doc, "  ☐ One master file in the agreed aspect ratio")
    body(doc, "  ☐ Branded intro     ☐ Branded outro     ☐ Thumbnail     ☐ Cover frame")
    body(doc, "  ☐ Subtitles / captions file       ☐ Cliffhanger opener")
    body(doc, "  ☐ YouTube kit (timestamps + optimized title + description)")
    field(doc, "Number of extra export versions", BLANK_SHORT)

    heading(doc, "4. Timeline", level=2)
    body(doc, "  ☐ Standard delivery: 7 calendar days from receipt of full footage and brief.")
    body(doc, "  ☐ Rush delivery: 5 calendar days from receipt of full footage and brief, +25%.")
    field(doc, "Agreed final delivery date", BLANK_MED)
    body(
        doc,
        "Delays caused by late footage, late brief responses, or scope "
        "changes from the Client extend the delivery date by the equivalent "
        "number of days.",
    )

    heading(doc, "5. Fees and payment terms", level=2)
    field(doc, "Total project fee (₱)", BLANK_MED)
    body(doc, "Payment is split as follows:")
    body(
        doc,
        "  · 50% of the total fee is due before work begins (down payment).",
    )
    body(
        doc,
        "  · 50% of the total fee is due before the midpoint deliverable "
        "(e.g. for a 4-video monthly retainer, before video 3 starts).",
    )
    body(
        doc,
        "Both halves are prepayments. I begin work only after each "
        "prepayment is received and confirmed.",
    )
    field(doc, "Down payment received on", BLANK_MED)
    field(doc, "Balance received on", BLANK_MED)

    heading(doc, "6. Revisions", level=2)
    body(
        doc,
        "One round of revisions is included in the project fee. Additional "
        "rounds are quoted separately. Revision requests must be submitted "
        "in writing within 7 days of delivery; later requests are treated "
        "as new work.",
    )

    heading(doc, "7. Footage and ownership", level=2)
    body(
        doc,
        "The Client warrants that they have the legal right to use, modify, "
        "and publish all footage, music, and assets supplied to me. The "
        "Client indemnifies me against any third-party claims arising from "
        "those rights.",
    )
    body(
        doc,
        "Upon final payment, the Client owns the final deliverables. I "
        "retain the right to display the work in my portfolio and case "
        "studies unless the Client requests otherwise in writing before "
        "delivery.",
    )
    body(
        doc,
        "Project files (raw cuts, source projects) remain mine and are "
        "kept for 90 days after final delivery, after which they may be "
        "deleted.",
    )

    heading(doc, "8. Confidentiality", level=2)
    body(
        doc,
        "I will treat all unpublished footage, scripts, and brand "
        "information as confidential and will not share them outside this "
        "engagement without the Client's written permission.",
    )

    heading(doc, "9. Cancellation", level=2)
    body(
        doc,
        "If the Client cancels after the down payment is received but "
        "before I have delivered the first cut, the down payment is "
        "non-refundable as compensation for booked time. If cancellation "
        "happens after the first cut is delivered, the full project fee "
        "is due.",
    )
    body(
        doc,
        "I may terminate this agreement and refund any unearned portion "
        "of the fee if the Client repeatedly delays beyond agreed "
        "timelines or asks for work outside the agreed scope without "
        "negotiation.",
    )

    heading(doc, "10. Liability", level=2)
    body(
        doc,
        "My total liability under this agreement is limited to the total "
        "project fee paid. I am not liable for indirect or consequential "
        "losses.",
    )

    heading(doc, "11. Governing law", level=2)
    body(
        doc,
        "This agreement is governed by the laws of the Republic of the "
        "Philippines.",
    )

    heading(doc, "12. Signatures", level=2)
    doc.add_paragraph()
    field(doc, "Editor signature", BLANK_LONG)
    field(doc, "Printed name", "Audrey Baliao")
    field(doc, "Date", BLANK_MED)
    doc.add_paragraph()
    field(doc, "Client signature", BLANK_LONG)
    field(doc, "Printed name", BLANK_LONG)
    field(doc, "Date", BLANK_MED)

    out = PUBLIC / "dhey-creates-service-contract.docx"
    doc.save(str(out))
    return out


def main() -> int:
    a = build_intake_form()
    b = build_service_contract()
    print(f"✓ Wrote {a.relative_to(ROOT)}")
    print(f"✓ Wrote {b.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
